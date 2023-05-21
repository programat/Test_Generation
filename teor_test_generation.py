import random

import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, Inches
from lxml import etree
import latex2mathml.converter


def printToMathml(paragraph, formula, con=True):
    if con:
        stri = latex2mathml.converter.convert(formula)
    else:
        stri = formula
    tree = etree.fromstring(stri)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    func = transform(tree)
    paragraph._element.append(func.getroot())


# все ответы будут записываться в формате:
# taskt = {40: [[«Чему равно математическое ожидание при экспоненциальном распределении с «,, «latex», «параметром»], [(«lals», 1)]]}

tasks = {1: (['Элементарное событие – это '], ['Единичный исход', 'Число', 'Эксперимент', 'Вывод']),
         2: (['Событие – это '],
             ['Подмножество множества элементарных событий', 'Утверждение', 'Пространство элементарных событий',
              'Доказательство']),
         3: (['Вероятность – это '],
             ['Степень возможности наступления некоторого события', 'Утверждение', 'Множество', 'Эксперимент']),
         4: (['Вероятность наступления некоторого события не может быть равна '], ['2', '1', '0', '0.5']),
         5: (['P(A+B)= (сложение вероятностей)'], ['P(A)+P(B)', 'P(A)-P(B)', 'P(AB)+P(A)', 'P(AB)+P(B)']),
         6: (['Случайное событие – это '],
             ['Может как произойти так и не произойти', 'Доказанное утверждение', 'Очевидное свойство',
              'Положительное число']),
         7: (['Случайная величина есть'], ['Функция элементарных событий', 'Число', 'Вывод ','Эксперимент']),
         8: (['Функция распределения случайной величины есть'], ['Функция одного действительного переменного', 'Функция элементарных событий', 'Функция многих действительных переменных','Функция двух действительных переменных']),
         9: (['Вероятность того, что непрерывная случайная величина примет конкретное значение равна'], ['0', '1', 'Зависит от задачи','Нет правильных ответов']),
         10: (['Какие значения может принимать функция распределения?'], [(r'0 \leq F(x) \leq 1', 1), (r'-\infty \leq F(x) \leq +\infty', 1), (r'F(x) > 0', 1), (r'F(x) \neq 1', 1)]),
         11: (['Что означает операция А+В?'], ['произошло хотя бы одно из двух событий А или В', 'событие А влечет за собой событие В', 'совместно осуществились события А и В', 'Событие В влечет за собой событие А']),
         12: (['Что означает операция АВ?'], ['Произошло и событие А, и событие В', 'Произошло хотя бы одно из двух событий А или В', 'Событие А влечет за собой событие В', 'Ни одно из событий не произошло']),
         13: (['Выберите неверное утверждение'], ['Вероятность появления одного из противоположных событий всегда больше вероятности другого', 'Сумма вероятностей двух противоположных событий равна единице', 'Если два события единственно возможны и несовместны, то они называются противоположными', 'Событие, которое никогда не произойдет, является невозможным']),
         14: ([(r'A', 1), ' и ', (r'B', 1),' - независимые события. Тогда справедливо следующее утверждение:'], [(r'P(B/A) = P(B)', 1), (r'P(A/B) = P(B)', 1), (r'P(A \cup B) = P(B)', 1), (r'p(A \cap B) = 0', 1)]),
         15: (['Равномерное распределение случайной величины имеет вид'], [(r'P_m = 1/n', 1), (r'P(X = m) = C^m_np^mq^{n-m}', 1), (r'P(X = m) = p^mq^{n-m}', 1), (r'P(X = m) = \frac{\lambda^m e^{-\lambda}}{m!}', 1)]),
         16: (['Распределение Пуассона случайной величины имеет вид'], [(r'P(X = m) = \frac{\lambda^m e^{-\lambda}}{m!}', 1), (r'P_m = 1/n', 1), (r'P(X = m) = p^mq^{n-m}', 1), (r'P(X = m) = C^m_np^mq^{n-m}',1)]),
         17: (['Биномиальное распределение случайной величины имеет вид'], [(r'P(X = m) = C^m_np^mq^{n-m}', 1), (r'P_m = 1/n', 1), (r'P(X = m) = p^mq^{n-m}', 1), (r'P(X = m) = \frac{\lambda^m e^{-\lambda}}{m!}',1)]),
         18: (['Распределение Бернулли случайной величины имеет вид'], [(r'P(X = m) = p^mq^{n-m}', 1), (r'P_m = 1/n', 1), (r'P(X = m) = C^m_np^mq^{n-m}', 1), (r'P(X = m) = \frac{\lambda^m e^{-\lambda}}{m!}',1)]),
         19: (['Как называется число ', (r'm_0', 1), 'наступления события в n независимых испытаниях, в каждом из которых вероятность появления события равна p, определяемое из неравенства ', (r'np-q\leq m_0\leq np+p', 1), '?'], ['Наивероятнейшее', 'Наибольшее', 'Оптимальное','Минимальное']),
         20: (['Максимальное значение произведения вероятностей противоположных событий равно'], ['0.25', '0.5', '1', '0.54']),
         21: (['Парный коэффициент корреляции r(X Y), изменяется в пределах'], [(r'-1\leq r(X,Y)\leq 1', 1), (r'0\leq r(X,Y)\leq 1', 1), (r'-\infty\leq r(X,Y)\leq +\infty', 1), (r'0\leq r(X,Y)\leq +\infty',1)]),
         22: (['Парный коэффициент корреляции равен –1. Это означает'], ['Отрицательную линейную связь', 'Наличие нелинейной функциональной связи', 'Отсутствие связи', 'Положительную линейную связь']),
         23: (['Вероятности появления заданного числа благоприятных исходов в схеме Бернулли описываются'], ['Биноминальным распределением', 'Геометрическим распределением', 'Равномерным распределением на отрезке', 'Однородным распределением']),
         24: (['Математического ожидания не существует у случайной величины'], ['Распределенной по Коши', 'Равномерно распределенной на отрезке', 'Имеющей нормальное распределение', 'Неравномерно распределенной на отрезке']),
         25: (['Закон больших чисел выводится из неравенства Чебышева при условии существования у случайной величины'], ['Конечного второго момента', 'Конечного математического ожидания', 'Плотности', 'Дисперсии']),
         26: (['Характеристическая функция случайной величины есть'], ['Комплекснозначная функция действительного переменного', 'Аналитическая функция комплексного переменного', 'Действительная функция комплексного переменного', 'Мнимая функция комплексного переменного']),
         27: (['Если характеристическая функция случайной величины имеет производную в точке нуль, то'], ['Случайная величина имеет конечное математическое ожидание', 'Случайная величина имеет плотность', 'Случайная величина имеет конечный момент второго порядка', 'Все варианты неверные']),
         28: (['Характеристическая функция нормального стандартного распределения равна'], [(r'e^{-\frac{r^2}{2}}', 1), (r'e^{it}', 1), (r'1', 1), (r'1 - e^{it}',1)]),
         29: (['Зная характеристическую функцию можно определить функцию распределения'],
              ['Произвольной случайной величины', 'Непрерывной случайной величины', 'Простой случайной величины',
               'Невозможно определить функцию распределения']),
         30: (['Определите закон распределения непрерывной случайной величины, если плотность распределения имеет вид ', (r'<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>p(x)</mi></mrow><mo>=</mo><mrow><mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" align="center 2"><mtr><mtd><mi>0,x</mi><mo>&#8950;</mo><mi>[a,b]</mi></mtd></mtr><mtr><mtd><mfrac><mi>1</mi><mrow><mi>a-b</mi></mrow></mfrac><mo><mchar name="InvisibleTimes"/></mo><mrow><mi>x</mi><mo>&#8712;</mo><mi>[a,b]</mi></mrow></mtd></mtr></mtable></mfenced></mrow></math>', 1)], ['Равномерное распределение', 'Экспоненциальное распределение', 'Нормальное распределение', 'Биномиальное распределение']),
         31: (['Определите закон распределения непрерывной случайной величины, если плотность распределения имеет вид ', (r'<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>p(x)</mi></mrow><mo>=</mo><mrow><mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" align="center 2"><mtr><mtd><mrow><mi>&#955;</mi><mo><mchar name="InvisibleTimes"/></mo><mrow><msup><mi>e</mi><mrow><mi>-&#955;x</mi></mrow></msup></mrow></mrow><mo><mchar name="InvisibleTimes"/></mo><mi>,x&#8805;0</mi></mtd></mtr><mtr><mtd><mrow><mi>0,x&lt;0</mi></mrow></mtd></mtr></mtable></mfenced></mrow></math>', 1)], ['Экспоненциальное распределение', 'Нормальное распределение', 'Равномерное распределение','Биномиальное распределение']),
         32: (['Определите закон распределения непрерывной случайной величины, если плотность распределения имеет вид  ', (r'p(x) = \frac{1}{\sigma \sqrt{2\pi}} e^{-\frac{1}{2}\frac{(x-a)^2}{\sigma^2}}', 1)], ['Нормальное распределение', 'Биномиальное распределение', 'Распределение Бернулли','Распределение Пуассона']),
         33: (['Выберете неверное утверждение:'], ['Функция распределения F(x, у) есть отрицательная функция, заключенная между нулем и единицей', 'Функция распределения F(x, у) есть неубывающая функция по каждому из аргументов', ['Если хотя бы один из аргументов обращается в ', (r'-\infty', 1), ' функция распределения F(x, у)  равна нулю'], ['Если оба аргумента равны ', (r'+\infty', 1), ' то функция распределения равна единице']]),
         34: (['Двумерная случайная величина называется непрерывной, если ее функция распределения-'], ['непрерывная, дифференцируемая по каждому из аргументов, и существует вторая смешанная производная', 'непрерывная, дифференцируемая по каждому из аргументов, и существует третья смешанная производная', 'непрерывная', 'Ни один вариант не является верным']),
         35: (['Плотность распределения вероятностей непрерывной двумерной случайной величины –это'], ['Вторая смешанная частная производная ее функции распределения', 'Сумма всех вероятностей', 'Постоянная величина', 'Все варианты верные']),
         36: (['Выберете верный вариант'], ['Вероятность попадания непрерывной двумерной величины (X, Y) в область D равна ', 'Вероятность попадания непрерывной двумерной величины (X, Y) в область D равна', 'Плотность вероятности двумерной случайной величины есть отрицательная функция', 'полный объем тела, ограниченного поверхностью распределения и плоскостью Оху, равен -1']),
         37: (['Математическое ожидание постоянной равно'], ['Этой постоянной', '1', '2', 'Нет верного варианта']),
         38: (['Для каких случайных величин справедливо свойство математического ожидания M (X + Y) = MX + MY'], ['И для зависимых, и для независимых', 'Только для зависимых', 'Только для независимых','Нет верного варианта']),
         39: (['Чему равно математическое ожидание при пуассоновском распределении с параметром ', (r'\lambda', 1), '?'], [(r'\lambda', 1), (r'\frac{a+b}{2}', 1), (r'\frac{1}{\lambda}', 1), 'Нет верного ответа']),
         40: (['Чему равно математическое ожидание при экспоненциальном распределении с параметром ', (r'\lambda', 1), '?'], [(r'\frac{1}{\lambda}', 1), (r'\frac{a+b}{2}', 1), (r'\lambda', 1), 'Нет верного ответа']),
         41: (['Чему равно математическое ожидание при экспоненциальном распределении с параметром ', (r'\lambda', 1), '?'], [(r'\frac{a+b}{2}', 1), (r'\lambda', 1), (r'\frac{1}{\lambda}', 1), 'Нет верного ответа']),
         42: (['Какой вероятности соответствует медиана?'], ['0.5', '1', '0.25','Нет верного ответа']),
         43: (['Вставьте пропуск.\nЕсли Х – непрерывная случайная величина, то мода – __________________ плотности распределения'], ['Точка локального максимума', 'Точка локального минимума', 'Несуществующая точка', 'Нет верного ответа']),
         44: (['Числом, равным математическому ожиданию квадрата отклонения случайной величины от её математического ожидания называют'], ['Дисперсию', 'Моду', 'Медиану', 'Квантиль']),
         45: (['Выберете неверное'], [(r'DX = MX^2-MX', 1), (r'DX = MX^2-M^2X', 1), (r'DX = M(X^2)-(MX)^2', 1), (r'DX=M(X-MX)^2',1)]),
         46: (['Среднее квадратическое отклонение случайной величины'], [(r'\sigma_x = \sqrt{DX}', 1), (r'\sigma_x = DX^2', 1), (r'\sigma_x = \frac{DX}{2}', 1), (r'\sigma_x = e^{DX}',1)]),
         47: (['D(X+Y)='], ['DX+DY', 'D(XY)', 'DX+DY-D(XY)','0']),
         48: (['Верно ли равенство: ', (r'D(X-Y)=D(X)+D(Y)', 1), '?'], ['Да', 'Нет', 'Зависит от задачи','Нет верного варианта']),
         49: (['Каково значение дисперсии при нормальном распределении?'], [(r'\sigma^2', 1), (r'\lambda', 1), (r'\frac{1}{\lambda^2}', 1), (r'\frac{(a-b)^2}{12}',1)]),
         50: (['Каково значение дисперсии при экспоненциальном распределении?'], [(r'\frac{1}{\lambda^2}', 1), (r'\frac{(a-b)^2}{12}', 1), (r'\sigma^2', 1), (r'\lambda', 1)], 1),
         51: (['Каково значение дисперсии при распределении Пуассона?'], [(r'\lambda', 1), (r'\frac{1}{\lambda^2}', 1), (r'\frac{(a-b)^2}{12}', 1), (r'\sigma^2',1)]),
         52: (['Каково значение дисперсии при равномерном распределении?'], [(r'\frac{(a-b)^2}{12}', 1), (r'\sigma^2', 1), (r'\lambda', 1), (r'\frac{1}{\lambda^2}',1)])
}



# задание 32 нужно смотреть отдельно, так как там формулы в вариантах ответов непосредственно

def generate_teor_tests(num_tests):
    document = docx.Document()  # документ с теоретическими тестами

    # задание стиля для header
    style_header = document.styles.add_style('f_header', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
    style_header.font.name = 'Times New Roman'
    style_header.font.size = docx.shared.Pt(14)

    # задание стиля для заданий
    style_task = document.styles.add_style('f_tasks', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
    style_task.font.name = 'Times New Roman'
    style_task.font.size = docx.shared.Pt(16)


    #Создаем документ для ответов
    doc_answers= docx.Document()
    title = doc_answers.add_paragraph(f'Ответы для тестов «Варианты (1-{num_tests})»')
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    font = title.style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(16)
    font.bold = True
    run = title.add_run()
    run.add_break(docx.enum.text.WD_BREAK.LINE)

    table = doc_answers.add_table(rows=1 + num_tests, cols=7, style='Table Grid')
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    row_cells = table.rows[0].cells
    row_cells[0].text = 'B\№'
    for i in range(1, 7):
        row_cells[i].text = str(i)
    answers_forTable = []



    for i in range(1, num_tests + 1):
        # добавление параграфа с вариантом
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Вариант (№{i})')
        run.style = style_header
        run.font.bold = True
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # добавление блока с заговоком
        paragraph = document.add_paragraph()
        run = paragraph.add_run('Тест по теме «Теория вероятностей и математическая статистика»\n')
        run.style = style_header
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        answers_row = []

        # Формируем массив заданий
        questions = []
        for j in range(0, 6):
            if j != 5:
                number = random.randint(1+j*9, 9+j*9)
            else:
                number = random.randint(46,52)
            questions.append(number)

        for j in range(0, 6):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f'{j+1}. ')
            run.style = style_task
            run.bold = True

            task = tasks[questions[j]][0]
            for item in task:
                print(item)
                if isinstance(item, tuple):
                    if questions[j]==30 or questions[j]==31:
                        printToMathml(paragraph, item[0], False)
                    else:
                        printToMathml(paragraph, item[0])
                else:
                    run = paragraph.add_run(item)
                    run.style = style_task

            #Выводим в документ
            task = tasks[questions[j]][1]
            realAnswer = task[0]
            answers = []
            for item in task:
                answers.append(item)
            random.shuffle(answers)
            answers_row.append(str(answers.index(realAnswer)))


            counter = 0
            for item in answers:
                print(item)
                paragraph = document.add_paragraph(f"{chr(ord('А')+counter)}) ")
                font = paragraph.style.font
                font.size = Pt(14)
                counter = counter+1
                if isinstance(item, tuple):
                    printToMathml(paragraph, item[0])
                elif len(item) == 3:
                    run=paragraph.add_run(item[0])
                    run.style = style_task
                    printToMathml(paragraph,item[1][0])
                    run=paragraph.add_run(item[2])
                    run.style = style_task
                else:
                    run = paragraph.add_run(item)
                    run.style = style_task

        p = document.add_paragraph('')
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        answers_forTable.append(answers_row)



    print(answers_forTable)
    for row in table.rows:
        for cell in row.cells:
            cell.width = docx.shared.Inches(0.5)
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # устанавливаем выравнивание ячеек по центру по вертикали
                for run in paragraph.runs:
                    run.font.bold = True  # убираем жирный шрифт

    for i in range(1, num_tests + 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)
        for j in range(1, 6 + 1):
            row_cells[j].text = chr(ord('А')+int(answers_forTable[i-1][j-1]))

# Выравниваем текст в ячейках по центру
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = docx.shared.Pt(14)
                    run.font.name = 'Times New Roman'
                    run.font.bold = False
    cells = table.column_cells(0)
    for cell in cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
    cells = table.row_cells(0)
    for cell in cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True

    document.save('text_teor.docx')
    doc_answers.save('answers.docx')

if __name__ == '__main__':
    generate_teor_tests(100)
