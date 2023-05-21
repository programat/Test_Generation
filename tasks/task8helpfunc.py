from fractions import Fraction
import random

import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, Inches
from lxml import etree

def CreateCosCoef(a):
    PossibleCoefs: list = []
    for i in range(1, int(a / 2) + 1):
        PossibleCoefs.append(i)
    PossibleCoefsOfPi = list(range(0, 11, 2))
    c = random.choice(PossibleCoefs)
    b = random.choice(PossibleCoefsOfPi) * a * 2
    return c + b


def CreateInterval():
    # Выбираем коэф для pi
    CoefSign: int = random.randint(0, 1)

    # Формирируем начало промежутка,где функция обращается в cos
    Picoef: list = list(range(0 + CoefSign, 10, 2))
    begincoef = random.choice(Picoef)
    beginInterval = begincoef

    # Формируем конечную точку промежутка
    possibleEnds: list = [Fraction(1, 6), Fraction(1, 4), Fraction(1, 3), Fraction(1, 2)]
    endInterval = random.choice(possibleEnds)

    return beginInterval, endInterval + begincoef


def CalcutionOfCoef():
    beginInterval, endInterval = CreateInterval()
    CosCoef: int = CreateCosCoef(endInterval.denominator)

    # Вычисляем точное значение sin(x)
    oldEndOfInterval = endInterval


    endInterval = Fraction(endInterval.numerator*CosCoef, endInterval.denominator)
    # print(endInterval)
    while endInterval > 2:
        endInterval = endInterval-2

    # print(endInterval)

    coef = 1*CosCoef
    if endInterval.numerator > endInterval.denominator:
        coef = -1*(CosCoef)

    a = endInterval.denominator
    match a:
        case 2:
            coef = coef*1
            s = '<mi>'+str(coef)+'</mi>'
        case 3:
            coef = coef * 2
            s = '<mfrac>'+'<mi>'+str(coef)+'</mi>'+'<mrow><msqrt><mi>3</mi></msqrt></mrow></mfrac>'
        case 4:
            coef = coef * 2
            s = '<mfrac>'+'<mi>'+ str(coef) +'</mi>'+'<mrow><msqrt><mi>2</mi></msqrt></mrow></mfrac>'
        case 6:
            coef = coef * 2
            s = '<mi>' + str(coef) + '</mi>'
        case defaul:
            return 0
    return '<math xmlns="http://www.w3.org/1998/Math/MathML" >' + s +'</math>', beginInterval, oldEndOfInterval, CosCoef

def printTask8(document):
    # Создаем нужные данные для функции распределения

    answerString, beginInterval, oldEndOfInterval, CosCoef = CalcutionOfCoef()

    # create MathML structure
    s0 = '<mrow><mi>f</mi><mo>&#x2061;</mo><mrow><mo>(</mo><mi>x</mi><mo>)</mo></mrow></mrow><mo>=</mo>'
    s1 = '<mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" align="center 2">'

    # Делаем mathml для начала интервала
    if beginInterval.numerator == 0:
        beginIntervalstr = '<mi>0</mi>'
    else:
        if beginInterval.numerator == 1:
            beginInterval = ''
        else:
            beginInterval = str(beginInterval.numerator)
        beginIntervalstr = '<mrow><mi>' + beginInterval + '</mi><mo><mchar name="InvisibleTimes"/></mo><mi>&#x03c0;</mi></mrow>'
    s2 = '<mtr><mtd><mrow><mi>0</mi><mo>,</mo><mspace width="20px"/><mrow><mi>при x</mi><mo>&#x2264;</mo><mrow>' + beginIntervalstr + '</mrow></mrow></mrow></mtd></mtr>'

    # Делаем косинус
    if CosCoef == 1:
        CosCoef = ''
    cosStr = '<mtr><mtd><mrow><mi>Ccos</mi><mo>&#x2061;</mo><mrow><mo>(</mo>' + '<mrow><mi>' + str(
        CosCoef) + '</mi><mo><mchar name="InvisibleTimes"/></mo><mi>x</mi></mrow>' + '<mo>)</mo></mrow></mrow>'

    # Делаем конец интервала
    if oldEndOfInterval.numerator == 1:
        endInt = ''
    else:
        endInt = str(oldEndOfInterval.numerator)
    endIntervalStr = '<mfrac><mrow><mstyle><mi>' + endInt + '</mi></mstyle><mo><mchar name="InvisibleTimes"/></mo><mi>&#x03c0;</mi></mrow>' + '<mi>' + str(
        oldEndOfInterval.denominator) + '</mi></mfrac>'

    s3 = cosStr + '<mo>,</mo><mrow><mspace width="5px"/><mi>при ' + beginIntervalstr + '</mi><mo><mchar name="InvisibleTimes"/></mo><mspace width="5px"/><mrow><mi></mi><mo>&#x003c;</mo><mrow><mi>x</mi><mo>&#x2264;</mo>' + endIntervalStr + '</mrow></mrow></mrow></mtd></mtr>'

    # Делаем последнюю строку
    s4 = '<mtr><mtd><mrow><mi>0</mi></mrow><mo>,</mo><mrow><mrow><mspace width="5px"/><mi>при x</mi><mspace width="5px"/></mrow><mo>&#x003e;</mo>' + endIntervalStr + '</mrow></mtd></mtr>'
    s5 = '</mtable></mfenced>'
    tree = etree.fromstring(
        '<math xmlns="http://www.w3.org/1998/Math/MathML">' + s0 + s1 + s2 + s3 + s4 + s5 + '</math>')
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    func = transform(tree)

    # Создаем другие варианты ответа и мешаем их
    answer = []
    tree = etree.fromstring(answerString)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    answerroot = transform(tree)
    answer.append(answerroot)

    if oldEndOfInterval.numerator == 1:
        old=random.randint(2,12)
    else:
        old=str(oldEndOfInterval.numerator)
    ans = '<mfrac><mrow><mi>' + str(oldEndOfInterval.denominator) + '</mi><mo><mchar name="InvisibleTimes"/></mo><mi>&#x03c0;</mi></mrow>' + '<mi>' + str(old) + '</mi></mfrac>'
    tree = etree.fromstring('<math xmlns="http://www.w3.org/1998/Math/MathML">' + ans + '</math>')
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    answer.append(new_dom)

    tree = etree.fromstring('<math xmlns="http://www.w3.org/1998/Math/MathML">' + beginIntervalstr + '</math>')
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    answer.append(new_dom)

    i = random.randint(2, 6)
    j = random.randint(i, 8)
    ans = '<mfrac><mrow><msqrt><mi>' + str(i) + '</mi></msqrt></mrow>' + '<mrow><msqrt><mi>' + str(
        j) + '</mi></msqrt><mo>-</mo><mi>' + str(random.randint(j, 10)) + '</mi></mrow>' + '</mfrac>'
    tree = etree.fromstring('<math xmlns="http://www.w3.org/1998/Math/MathML">' + ans + '</math>')
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    answer.append(new_dom)

    p_find = answer[0]
    random.shuffle(answer)

    # Вставляем в документ
    task = "Непрерывная случайная величина X задана плотностью распределения вероятностей:\t"
    p = document.add_paragraph(task, style='List Number')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(16)

    run = p.add_run()
    run.add_break()
    run.add_break(WD_BREAK.LINE)
    p._element.append(func.getroot())

    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run('Тогда параметр C принимает значение:\t')
    font = run.font
    font.size = Pt(16)
    run.add_break()

    table = document.add_table(rows=1, cols=4)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    row_cells = table.rows[0].cells

    for i in range(0, 4):
        s = chr(ord('а') + i) + ') '
        m = row_cells[i].add_paragraph(s)
        m.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        m._element.append(answer[i].getroot())

    for cells in row_cells:
        for par in cells.paragraphs:
            if len(par.text) == 0:
                p = par._element
                p.getparent().remove(p)
                p._p = p._element = None

    return chr(answer.index(p_find) + 1072).capitalize()
