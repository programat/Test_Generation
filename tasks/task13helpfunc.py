from fractions import Fraction
import random

import lxml
import sympy
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, Inches
from lxml import etree
from sympy import symbols, mathml
import math

def printTask13(document, M, D):

    #Формирируем ответы
    s1 = '<math xmlns="http://www.w3.org/1998/Math/MathML" display="block"><mfrac><mi>1</mi><mrow><mi>'
    s2 = '</mi><mo><mchar name="InvisibleTimes"/></mo><msqrt><mi>2</mi><mo><mchar name="InvisibleTimes"/></mo><mi>&#x03c0;</mi></msqrt></mrow></mfrac> <mo><mchar name="InvisibleTimes"/></mo>'
    answers = []
    x, e = symbols('x e')

    expr = e**((x-M)**2/(2*D))
    exprxml = mathml(expr, printer='presentation')
    dstr = str(int(math.sqrt(D)))
    if D == 1:
        dstr = ''
    formulastr = s1+dstr+s2+exprxml+'</math> '
    ans1='<math xmlns="http://www.w3.org/1998/Math/MathML">'+formulastr+'</math>'
    tree = etree.fromstring(ans1)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    answerroot = transform(tree)
    answers.append(answerroot)

    expr = e ** ((x - D) ** 2 / (2 * D))
    exprxml = mathml(expr, printer='presentation')
    mstr = str(M)
    if M == 1:
        mstr = ''
    formulastr = s1 + mstr + s2 + exprxml + '</math> '
    ans1 = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + formulastr + '</math>'
    tree = etree.fromstring(ans1)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    answerroot = transform(tree)
    answers.append(answerroot)

    if D == 1:
        dstr = ''
    expr = e ** ((x - M) ** 2 / (2 * M))
    exprxml = mathml(expr, printer='presentation')
    formulastr = s1 + str(int(math.sqrt(D))) + s2 + exprxml + '</math> '
    ans1 = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + formulastr + '</math>'
    tree = etree.fromstring(ans1)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    answerroot = transform(tree)
    answers.append(answerroot)


    expr = e ** ((x - D) ** 2 / (2 * M))
    exprxml = mathml(expr, printer='presentation')
    if D!=1:
        formulastr = s1 + str(D) + s2 + exprxml + '</math> '
    else:
        formulastr = s1 + '' + s2 + exprxml + '</math> '
    ans1 = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + formulastr + '</math>'
    tree = etree.fromstring(ans1)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    answerroot = transform(tree)
    answers.append(answerroot)

    random.shuffle(answers)

    table = document.add_table(rows=2, cols=2)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    ch=0
    for j in range(0,2):
        row_cells = table.rows[j].cells
        for i in range(0, 2):
            s = chr(ord('а') + ch) + ')'
            m = row_cells[i].add_paragraph(s)
            m.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            m._element.append(answers[ch].getroot())
            ch=ch+1
    for strings in table.rows:
        for cells in strings.cells:
            for par in cells.paragraphs:
                if len(par.text) == 0:
                    p = par._element
                    p.getparent().remove(p)
                    p._p = p._element = None
