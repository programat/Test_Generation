from fractions import Fraction
import random

import lxml
import sympy
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, Inches
from lxml import etree
from sympy import symbols, mathml
import math


def createInterval():
    begin = 0
    end = random.randint(1, 7)
    return begin, end


def printTask12(document):
    beginInerval, endInterval = createInterval()
    power = random.randint(2, 6)
    fractionCoef = Fraction(power, endInterval**power)

    #Считаем дисперсию
    dispersion = fractionCoef*endInterval**(power+2)/(power+2)-(fractionCoef*endInterval**(power+1)/(power+1))**2
    print(dispersion)
    print(endInterval,power)
    print(str(fractionCoef)+"x^"+str(power-1))

    sfunc = '<mrow><mi>f</mi><mo>&#x2061;</mo><mrow><mo>(</mo><mi>x</mi><mo>)</mo></mrow></mrow><mo>=</mo>'
    s1 = '<mtr><mtd><mrow><mi>0</mi><mo>,</mo><mrow><mi>при x</mi><mo>&#x2264;</mo><mn>0</mn></mrow></mrow></mtd></mtr>'

    x = symbols('x')
    fractionCoef = Fraction(power, endInterval ** (power))
    expr = x ** (power - 1)
    exprxml = mathml(expr, printer='presentation')

    if (fractionCoef.numerator == 1):
        strFracNum = ''
    else:
        strFracNum = str(fractionCoef.numerator)

    s2 = '<mtr><mtd><mrow><mfrac><mrow>' + '<mi>' + strFracNum + '</mi>' + '<mo><mchar name="InvisibleTimes"/></mo>' + '<mrow>' + exprxml + '</mrow></mrow>' + '<mi>' + str(
        fractionCoef.denominator) + '</mi></mfrac></mrow><mo>,</mo><mrow><mi>при 0</mi><mo>&#x003c;</mo><mrow><mi>x</mi><mo>&#x2264;</mo><mi>' + str(
        endInterval) + '</mi></mrow></mrow></mtd></mtr>'

    s3 = '<mtr><mtd><mrow><mi>0</mi></mrow><mo>,</mo><mrow><mrow><mi>при x</mi></mrow><mo>&#x003e;</mo><mi>' + str(
        endInterval) + '</mi></mrow></mtd></mtr>'

    resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + sfunc + '<mrow><mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" >' + s1 + s2 + s3 + '</mtable></mfenced></mrow>' + '</math>'
    tree = etree.fromstring(resultString)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    result = transform(tree)

    # Выводим условие
    p = document.add_paragraph('Непрерывная случайная величина X задана плотностью распределения вероятностей:\t',
                               style='List Number')
    run = p.add_run()
    run.add_break()
    run.add_break(WD_BREAK.LINE)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(16)
    p._element.append(result.getroot())


    # Составляем ответы
    answers: list = []

    if(dispersion.denominator==1):
        strDispersionDenominator='2'
    else:
        strDispersionDenominator=str(dispersion.denominator)
    resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' +'<mfrac><mi>'+str(dispersion.numerator)+'</mi><mi>'+strDispersionDenominator + '</mi></mfrac></math>'
    tree = etree.fromstring(resultString)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    result = transform(tree)
    answers.append(result)

    if dispersion.numerator==1:
        strDispersionNumerator='2'
    else:
        strDispersionNumerator=str(dispersion.numerator)
    resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + '<mfrac><mi>' + str(
        dispersion.denominator) + '</mi><mi>' + strDispersionNumerator + '</mi></mfrac></math>'
    tree = etree.fromstring(resultString)
    transform = etree.XSLT(xslt)
    result = transform(tree)
    answers.append(result)

    if (dispersion.denominator == 1):
        strDispersionDenominator = '2'
    else:
        strDispersionDenominator = str(dispersion.denominator+dispersion.numerator)
    resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + '<mfrac><mi>' + str(
        dispersion.numerator) + '</mi><mi>' + strDispersionDenominator + '</mi></mfrac></math>'
    tree = etree.fromstring(resultString)
    transform = etree.XSLT(xslt)
    result = transform(tree)
    answers.append(result)

    if (dispersion.denominator == 1):
        strDispersionDenominator = '2'
    else:
        strDispersionDenominator = str(dispersion.denominator)
    resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + '<mfrac><mi>' + str(
        dispersion.numerator+dispersion.denominator) + '</mi><mi>' + strDispersionDenominator + '</mi></mfrac></math>'
    tree = etree.fromstring(resultString)
    transform = etree.XSLT(xslt)
    result = transform(tree)
    answers.append(result)

    random.shuffle(answers)


    p = document.add_paragraph('')
    run = p.add_run('Тогда ее дисперсия равна:\t')
    font = run.font
    font.size = Pt(16)
    run.add_break()

    table = document.add_table(rows=1, cols=4)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    row_cells = table.rows[0].cells
    for i in range(0, 4):
        s = chr(ord('а') + i) + ') '
        m = row_cells[i].add_paragraph(s)
        m.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        m.style.font.name = 'Times New Roman'
        m.style.font.size = Pt(16)
        m._element.append(answers[i].getroot())

    for cells in row_cells:
        for par in cells.paragraphs:
            if len(par.text) == 0:
                p = par._element
                p.getparent().remove(p)
                p._p = p._element = None