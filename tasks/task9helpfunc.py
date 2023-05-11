from fractions import Fraction
import random
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, Inches
from lxml import etree
from sympy import symbols, mathml


def createInterval():
        a: int = 0
        b: int = random.randint(a+2, 7)
        return a, b

def createInnerInterval(a,b):
    begin = random.randint(a, b-1)
    end = random.randint(begin+1, b)
    print(f"Начало внешнего интервала:{a};Конец внешнего интервала:{b}")
    print(f"Начало внутреннего интервала:{begin};Конец внутреннего интервала:{end}")
    return begin, end


def createFractCoef(beginint,endint):
    powerCoef: int = random.randint(2, 5)
    frac = Fraction(powerCoef, endint**powerCoef)
    print(f"Дробный коэф фукнции {frac};Степень {powerCoef}")
    return powerCoef-1, frac


def printTask9(document):
    beginint, endint = createInterval()
    beginInner, endInner = createInnerInterval(beginint, endint)
    powerCoef, frec = createFractCoef(beginint, endint)

    taskstr = 'Непрерывная случайная величина X задана плотностью распределения вероятностей:'
    p = document.add_paragraph(taskstr, style='List Number')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    s0 = '<mrow><mi>f</mi><mo>&#x2061;</mo><mrow><mo>(</mo><mi>x</mi><mo>)</mo></mrow></mrow><mo>=</mo>'
    s1 = '<mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" align="center 2">'
    s2 = '</mtable></mfenced>'

    firstfuncpart='<mtr><mtd><mspace width="10px"/><mrow><mi>0</mi><mo><mchar name="InvisibleTimes"/></mo><mrow><mi>,при x </mi><mo>&#x2264;</mo><mi>0</mi></mrow></mrow></mtd></mtr>'

    if frec.numerator == 1:
        coef = ''
    else:
        coef = str(frec.numerator)

    x = symbols('x')
    expr = x**powerCoef
    exprxml = mathml(expr, printer='presentation')
    secondpartfunc = '<mtr><mtd><mfrac><mrow><mi>'+coef+'</mi><mo><mchar name="InvisibleTimes"/></mo>'+exprxml+'</mrow><mi>'+str(frec.denominator)+'</mi></mfrac><mrow><mi>,при 0</mi><mo><mchar name="InvisibleTimes"/></mo><mrow><mrow><mo>&#x2264;</mo><mi>x</mi></mrow><mo>&#x2264;</mo><mi>'+str(endint)+'</mi></mrow></mrow></mtd></mtr>'

    coef=str(endint)
    thirdfuncpart='<mtr><mtd><mi>0 </mi><mo><mchar name="InvisibleTimes"/></mo><mrow><mi>,при x </mi><mo>&#x003e;</mo><mi>'+coef+'</mi></mrow></mtd></mtr>'

    s = '<math xmlns="http://www.w3.org/1998/Math/MathML" >' + s0+s1+firstfuncpart+secondpartfunc+thirdfuncpart+s2 +'</math>'

    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    func = transform(etree.fromstring(s))
    p = document.add_paragraph()
    p._element.append(func.getroot())


    #Формируем ответ
    rightAnswer= Fraction(beginInner, endInner)
    p = document.add_paragraph('Тогда вероятность ')
    func=''
    s = '<math xmlns="http://www.w3.org/1998/Math/MathML" >' + '<mrow><mi>P(</mi></mrow><mo><mchar name="InvisibleTimes"/></mo><mrow><mrow><mi>'+str(beginInner)+'</mi></mrow><mo>&#x2264;</mo><mrow><mi>X</mi><mo>&#x2264;</mo><mi>'+str(endInner)+')</mi></mrow></mrow>' + '</math>'
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    func = transform(etree.fromstring(s))
    p._element.append(func.getroot())
    run = p.add_run(' равна:\n')




    #Формириуем ответы
    answers: list = []
    ansFraction=(frec/(powerCoef+1))*((endInner)**(powerCoef+1)-(beginInner)**(powerCoef+1))
    print(f'Коэф функции распределения {frec*1/(powerCoef+1)}')
    print(f'Ответ {ansFraction}')
    s = '<math xmlns="http://www.w3.org/1998/Math/MathML" >' +'<mfrac><mi>'+str(ansFraction.numerator)+'</mi><mi>'+str(ansFraction.denominator) +'</mi></mfrac></math>'
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    func = transform(etree.fromstring(s))
    answers.append(func)

    s = '<math xmlns="http://www.w3.org/1998/Math/MathML" >' + '<mfrac><mi>' + str(
        random.randint(ansFraction.numerator+1, ansFraction.denominator)) + '</mi><mi>' + str(ansFraction.denominator) + '</mi></mfrac></math>'
    transform = etree.XSLT(xslt)
    func = transform(etree.fromstring(s))
    answers.append(func)

    s = '<math xmlns="http://www.w3.org/1998/Math/MathML" >' + '<mfrac><mi>' + str(1) + '</mi><mi>' + str(
        ansFraction.denominator+ansFraction.numerator) + '</mi></mfrac></math>'
    transform = etree.XSLT(xslt)
    func = transform(etree.fromstring(s))
    answers.append(func)

    s = '<math xmlns="http://www.w3.org/1998/Math/MathML" >' + '<mfrac><mi>' + str(ansFraction.numerator + 1) + '</mi><mi>' + str(
        ansFraction.denominator*2) + '</mi></mfrac></math>'
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    func = transform(etree.fromstring(s))
    answers.append(func)

    random.shuffle(answers)

    table = document.add_table(rows=1, cols=4)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    row_cells = table.rows[0].cells

    for i in range(0, 4):
        s = chr(ord('а') + i) + ')'
        m = row_cells[i].add_paragraph(s)
        m.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        m._element.append(answers[i].getroot())

    for cells in row_cells:
        for par in cells.paragraphs:
            if len(par.text) == 0:
                p = par._element
                p.getparent().remove(p)
                p._p = p._element = None


