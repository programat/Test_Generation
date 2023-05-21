from fractions import Fraction
import random

import lxml
import sympy
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, Inches
from lxml import etree
from sympy import symbols, mathml
import math


def createInteval():
    begin = 0
    end = random.randint(2, 10)
    return begin, end


def createPower():
    return random.randint(2, 5)


def printTask10(document):
	power = createPower()
	beginInterval, endInteval = createInteval()
	fractionCoef = Fraction(1, endInteval**(power))

	# Формируем вывод

	# Формируем условия
	sfunc = '<mrow><mi>F</mi><mo>&#x2061;</mo><mrow><mo>(</mo><mi>x</mi><mo>)</mo></mrow></mrow><mo>=</mo>'
	s1 = '<mtr><mtd><mrow><mi>0</mi><mo>,</mo><mrow><mi>при x</mi><mo>&#x2264;</mo><mn>0</mn></mrow></mrow></mtd></mtr>'

	x = symbols('x')
	expr = x**(power)
	exprxml = mathml(expr, printer='presentation')

	s2 = '<mtr><mtd><mrow><mfrac><mrow>'+'<mi>'+'</mi>'+'<mo><mchar name="InvisibleTimes"/></mo>'+'<mrow>'+exprxml+'</mrow></mrow>'+'<mi>'+str(fractionCoef.denominator)+'</mi></mfrac></mrow><mo>,</mo><mrow><mi>при 0</mi><mo>&#x003c;</mo><mrow><mi>x</mi><mo>&#x2264;</mo><mi>'+str(endInteval)+'</mi></mrow></mrow></mtd></mtr>'

	s3 = '<mtr><mtd><mrow><mi>1</mi></mrow><mo>,</mo><mrow><mrow><mi>при x</mi></mrow><mo>&#x003e;</mo><mi>'+str(endInteval)+'</mi></mrow></mtd></mtr>'

	resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' +sfunc+'<mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" >'+s1+s2+s3+'</mtable></mfenced>'+ '</math>'
	tree = etree.fromstring(resultString)
	xslt = etree.parse('MML2OMML.XSL')
	transform = etree.XSLT(xslt)
	result = transform(tree)

	#Выводим условие
	p = document.add_paragraph()
	# run = p.add_run()
	# run.add_break()
	# run.add_break(WD_BREAK.LINE)
	p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
	p.style.font.name = 'Times New Roman'
	p.style.font.size = Pt(16)
	p._element.append(result.getroot())

	p = document.add_paragraph()
	run = p.add_run('Тогда ее плотность распределения вероятностей имеет вид:\t')
	font=run.font
	font.size=Pt(16)
	run.add_break()
	p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

	# Начинаем делать ответы
	answers = []

	# Верный ответ
	sfunc = '<mrow><mi>f</mi><mo>&#x2061;</mo><mrow><mo>(</mo><mi>x</mi><mo>)</mo></mrow></mrow><mo>=</mo>'
	s1 = '<mtr><mtd><mrow><mi>0</mi><mo>,</mo><mrow><mi>при x</mi><mo>&#x2264;</mo><mn>0</mn></mrow></mrow></mtd></mtr>'

	x = symbols('x')
	fractionCoef = Fraction(power, endInteval ** (power))
	expr = x ** (power-1)
	exprxml = mathml(expr, printer='presentation')

	if(fractionCoef.numerator==1):
		strFracNum=''
	else:
		strFracNum=str(fractionCoef.numerator)

	s2 = '<mtr><mtd><mrow><mfrac><mrow>' + '<mi>' +strFracNum+ '</mi>' + '<mo><mchar name="InvisibleTimes"/></mo>' + '<mrow>' + exprxml + '</mrow></mrow>' + '<mi>' + str(
		fractionCoef.denominator) + '</mi></mfrac></mrow><mo>,</mo><mrow><mi>при 0</mi><mo>&#x003c;</mo><mrow><mi>x</mi><mo>&#x2264;</mo><mi>' + str(
		endInteval) + '</mi></mrow></mrow></mtd></mtr>'

	s3 = '<mtr><mtd><mrow><mi>0</mi></mrow><mo>,</mo><mrow><mrow><mi>при x</mi></mrow><mo>&#x003e;</mo><mi>' + str(
		endInteval) + '</mi></mrow></mtd></mtr>'

	resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + sfunc + '<mrow><mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" >' + s1 + s2 + s3 + '</mtable></mfenced></mrow>' + '</math>'
	tree = etree.fromstring(resultString)
	xslt = etree.parse('MML2OMML.XSL')
	transform = etree.XSLT(xslt)
	result = transform(tree)
	answers.append(result)

	#Остальные ответы

	sfunc = '<mrow><mi>f</mi><mo>&#x2061;</mo><mrow><mo>(</mo><mi>x</mi><mo>)</mo></mrow></mrow><mo>=</mo>'
	s1 = '<mtr><mtd><mrow><mi>0</mi><mo>,</mo><mrow><mi>при x</mi><mo>&#x2264;</mo><mn>0</mn></mrow></mrow></mtd></mtr>'

	x = symbols('x')
	fractionCoef = Fraction(1, endInteval ** (power))
	expr = x ** (power)
	exprxml = mathml(expr, printer='presentation')

	if (fractionCoef.numerator == 1):
		strFracNum = ''
	else:
		strFracNum = str(fractionCoef.numerator)

	s2 = '<mtr><mtd><mrow><mfrac><mrow>' + '<mi>' + strFracNum + '</mi>' + '<mo><mchar name="InvisibleTimes"/></mo>' + '<mrow>' + exprxml + '</mrow></mrow>' + '<mi>' + str(
		fractionCoef.denominator) + '</mi></mfrac></mrow><mo>,</mo><mrow><mi>при 0</mi><mo>&#x003c;</mo><mrow><mi>x</mi><mo>&#x2264;</mo><mi>' + str(
		endInteval) + '</mi></mrow></mrow></mtd></mtr>'

	s3 = '<mtr><mtd><mrow><mi>0</mi></mrow><mo>,</mo><mrow><mrow><mi>при x</mi></mrow><mo>&#x003e;</mo><mi>' + str(
		endInteval) + '</mi></mrow></mtd></mtr>'

	resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + sfunc + '<mrow><mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" >' + s1 + s2 + s3 + '</mtable></mfenced></mrow>' + '</math>'
	tree = etree.fromstring(resultString)
	xslt = etree.parse('MML2OMML.XSL')
	transform = etree.XSLT(xslt)
	result = transform(tree)
	answers.append(result)


	sfunc = '<mrow><mi>f</mi><mo>&#x2061;</mo><mrow><mo>(</mo><mi>x</mi><mo>)</mo></mrow></mrow><mo>=</mo>'
	s1 = '<mtr><mtd><mrow><mi>0</mi><mo>,</mo><mrow><mi>при x</mi><mo>&#x2264;</mo><mn>0</mn></mrow></mrow></mtd></mtr>'

	x = symbols('x')
	fractionCoef = Fraction(1, endInteval ** (power-1))
	expr = x ** (power)
	exprxml = mathml(expr, printer='presentation')

	if (fractionCoef.numerator == 1):
		strFracNum = ''
	else:
		strFracNum = str(fractionCoef.numerator)

	s2 = '<mtr><mtd><mrow><mfrac><mrow>' + '<mi>' + strFracNum + '</mi>' + '<mo><mchar name="InvisibleTimes"/></mo>' + '<mrow>' + exprxml + '</mrow></mrow>' + '<mi>' + str(
		fractionCoef.denominator) + '</mi></mfrac></mrow><mo>,</mo><mrow><mi>при 0</mi><mo>&#x003c;</mo><mrow><mi>x</mi><mo>&#x2264;</mo><mi>' + str(
		endInteval) + '</mi></mrow></mrow></mtd></mtr>'

	s3 = '<mtr><mtd><mrow><mi>0</mi></mrow><mo>,</mo><mrow><mrow><mi>при x</mi></mrow><mo>&#x003e;</mo><mi>' + str(
		endInteval) + '</mi></mrow></mtd></mtr>'

	resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + sfunc + '<mrow><mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" >' + s1 + s2 + s3 + '</mtable></mfenced></mrow>' + '</math>'
	tree = etree.fromstring(resultString)
	xslt = etree.parse('MML2OMML.XSL')
	transform = etree.XSLT(xslt)
	result = transform(tree)
	answers.append(result)



	sfunc = '<mrow><mi>f</mi><mo>&#x2061;</mo><mrow><mo>(</mo><mi>x</mi><mo>)</mo></mrow></mrow><mo>=</mo>'
	s1 = '<mtr><mtd><mrow><mi>0</mi><mo>,</mo><mrow><mi>при x</mi><mo>&#x2264;</mo><mn>0</mn></mrow></mrow></mtd></mtr>'

	x = symbols('x')
	fractionCoef = Fraction(power, endInteval ** (power))
	expr = x ** (power-1)
	exprxml = mathml(expr, printer='presentation')

	if (fractionCoef.numerator == 1):
		strFracNum = ''
	else:
		strFracNum = str(fractionCoef.numerator)

	s2 = '<mtr><mtd><mrow><mfrac><mrow>' + '<mi>' + strFracNum + '</mi>' + '<mo><mchar name="InvisibleTimes"/></mo>' + '<mrow>' + exprxml + '</mrow></mrow>' + '<mi>' + str(
		fractionCoef.denominator) + '</mi></mfrac></mrow><mo>,</mo><mrow><mi>при 0</mi><mo>&#x003c;</mo><mrow><mi>x</mi><mo>&#x2264;</mo><mi>' + str(
		endInteval) + '</mi></mrow></mrow></mtd></mtr>'

	s3 = '<mtr><mtd><mrow><mi>1</mi></mrow><mo>,</mo><mrow><mrow><mi>при x</mi></mrow><mo>&#x003e;</mo><mi>' + str(
		endInteval) + '</mi></mrow></mtd></mtr>'

	resultString = '<math xmlns="http://www.w3.org/1998/Math/MathML">' + sfunc + '<mrow><mfenced open="{" close="" separators=";;,"><mtable frame="solid" rowlines="solid" columnlines="solid" >' + s1 + s2 + s3 + '</mtable></mfenced></mrow>' + '</math>'
	tree = etree.fromstring(resultString)
	xslt = etree.parse('MML2OMML.XSL')
	transform = etree.XSLT(xslt)
	result = transform(tree)
	answers.append(result)

	p_find = answers[0]
	random.shuffle(answers)

	table = document.add_table(rows=2, cols=2)
	ch = 0
	for j in range(0, 2):
		row_cells = table.rows[j].cells
		for i in range(0, 2):
			s = chr(ord('а') + ch) + ') '
			m = row_cells[i].add_paragraph(s)
			m.style.font.name = 'Times New Roman'
			m.style.font.size = Pt(14)
			m._element.append(answers[ch].getroot())
			ch = ch + 1
	for strings in table.rows:
		for cells in strings.cells:
			for par in cells.paragraphs:
				if len(par.text) == 0:
					p = par._element
					p.getparent().remove(p)
					p._p = p._element = None

	return chr(answers.index(p_find) + 1072).capitalize()