# библиотеки для интерфейса
import tkinter as tk
from tkinter import ttk, messagebox
import os

# библиотеки для работы с word
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# библиотеки для генерации
from itertools import product
import random
import fractions as fr

# подключение внешних файлов
from tasks.task8helpfunc import printTask8
from tasks.task9helpfunc import printTask9
from tasks.task10helpfunc import printTask10
from tasks.task12helpfunc import printTask12
from tasks.task13helpfunc import printTask13
from teor_test_generation import printToMathml

from teor_test_generation import generate_teor_tests

tasks = {
    1: 'Игральная кость бросается три раза. Тогда вероятность того, что сумма выпавших очков не меньше шестнадцати, равна: ',
    2: 'Вероятность, что наудачу брошенная в круг точка окажется внутри вписанного в него квадрата равна:',
    3: 'Сантехник обслуживает три дома. Вероятность того, что в течение часа потребуется его помощь в первом доме, равна 0,15; во втором – 0,25; в третьем – 0,2. Тогда вероятность  того, что в течение часа потребуется его помощь хотя бы в одном доме, равна (приближенно равна):',
    4: 'Предприятие выплачивает 44 % всех зарплат разнорабочим, а 56 % – остальным. Вероятность того, что разнорабочий не получит зарплату в срок, равна 0,2; а для остальных эта вероятность составляет 0,1. Тогда вероятность того, что очередная зарплата будет выдана в срок, равна:',
    5: 'Имеются четыре коробки, в которых сидят по 3 белых и по 7 черных котят, и шесть коробок, в которых сидят по 8 белых и по 2 черных котенка. Из наудачу взятой коробки вынимается один котенок, который оказался белым. Тогда вероятность того, что этого котенка достали из первой серии коробок, равна:',
    6: ('Дискретная случайная величина X задана законом распределения вероятностей:', 'Тогда вероятность P(<) равна:'),
    7: ('Дискретная случайная величина X задана законом распределения вероятностей:', 'И вероятность P(<) = 0,6. Тогда значения a, b и c могут быть равны:'),
    11: ('Математическое ожидание дискретной случайной величины X,', 'заданной законом распределения вероятностей, равно 4,7. Тогда значение вероятности p2 равно:')
}

answers = dict()

def t1():
    n, r = (6, 3)
    items = range(1, n + 1)
    arrangements = list(product(items, repeat=r))

    # numbers = {10: 'десяти', 11: 'одиннадцати', 12: 'двенадцати', 13: 'тринадцати', 14: 'четырнадцати', 15: 'пятнадцати', 16: 'шестнадцати'}
    numbers = {4: 'четырех', 5: 'пяти', 6: 'шести', 7: 'семи'}
    choose = random.randrange(4, 7)
    task = tasks[1].replace('шестнадцати', numbers[choose], 1)
    # print(task)

    counter_target = 0
    counter_wrong1, counter_wrong2 = 0, 0
    for i in range(len(arrangements)):
        if sum(arrangements[i]) >= choose: counter_target += 1
        if choose != 16:
            if sum(arrangements[i]) >= choose + 1: counter_wrong1 += 1
            if sum(arrangements[i]) >= choose - 2: counter_wrong2 += 1
        else:
            if sum(arrangements[i]) >= choose - 1: counter_wrong1 += 1
            if sum(arrangements[i]) >= choose - 2: counter_wrong2 += 1

    p = fr.Fraction(counter_target, len(arrangements))
    w1 = fr.Fraction(counter_target - 2, len(arrangements))
    w2 = fr.Fraction(counter_wrong1, len(arrangements))
    w3 = fr.Fraction(counter_wrong2, len(arrangements))

    # print(P)

    return [task, p, w1, w2, w3]

def t3():
    flag = 0
    while not flag:
        flag = 1
        # заменить на одну строчку
        items = []
        i = 0.1
        while i <= .5:
            items.append(round(i, 2))
            i += 0.1
        numbers = [random.choice(items) for i in range(3)]
        j = random.randint(0, 2)

        numbers[j] += 0.05
        numbers[j] = round(numbers[j], 3)
        task = (tasks[3].replace('0,15', str(numbers[0]), 1).replace('0,25', str(numbers[1]), 1).replace('0,2', str(
            numbers[2]), 1)).replace('0.', '0,')
        # print(numbers)
        numbers_q = [1 - i for i in numbers]
        p = round(1 - numbers_q[0] * numbers_q[1] * numbers_q[2], 3)
        w1 = round(1 - numbers[0] * numbers[1] * numbers[2], 3)
        w2 = round(numbers_q[0] * numbers_q[1] * numbers_q[0], 3)
        w3 = round(numbers[0] * numbers[1] * numbers[2], 3)

        ans = [task, p, w1, w2, w3]
        # print('iteration', ans[1:])
        for i in range(1, len(ans)):
            for j in range(i + 1, len(ans)):
                if abs(ans[i] - ans[j]) <= 0.03 or ans[i] > 0.95 or ans[i] < 0.1 or ans[i] == ans[j]:
                    # print(ans[i], ans[j])
                    flag = 0
                    break
            if not flag:
                break
    # print('finally')
    # print(ans)
    return ans

def t4():
    flag = 0
    while not flag:
        flag = 1

        p_b1 = random.randint(30, 60)  # разнорабочих в процентах
        p_b2 = 100 - p_b1  # остальные
        p_b1_a, p_b2_a = 0, 0
        while p_b1_a == p_b2_a:
            p_b1_a = round(round(random.uniform(.1, .3) / .1) * .1, 5)  # вер-ть, что разнорабочий НЕ получит зп
            p_b2_a = round(round(random.uniform(.1, .3) / .1) * .1, 5)  # вер-ть, что обычный рабочий НЕ получит зп
            q_b1_a = 1 - round(p_b1_a, 5)
            q_b2_a = 1 - round(p_b2_a, 5)

        task = tasks[4].replace('44 %', '$1$', 1).replace('56 %', '$2$', 1).replace('0,2', '$3$', 1).replace('0,1',
                                                                                                             '$4$', 1)
        task = task.replace('$1$', f'{p_b1} %', 1).replace('$2$', f'{p_b2} %', 1).replace('$3$', f'{p_b1_a}',
                                                                                          1).replace('$4$', f'{p_b2_a}',
                                                                                                     1).replace('0.', '0,')

        p_b1 = round(p_b1 / 100, 5)
        p_b2 = round(p_b2 / 100, 5)

        p = round(p_b1 * q_b1_a + p_b2 * q_b2_a, 7)
        w1 = round(p_b1 * p_b1_a + p_b2 * p_b2_a, 7)
        w2 = round(p + .12, 7)
        w3 = round(w1 - .12, 7)

        ans = [task, p, w1, w2, w3]
        # print('iteration', ans)
        for i in range(1, len(ans)):
            for j in range(i + 1, len(ans)):
                if abs(ans[i] - ans[j]) <= 0.01 or ans[i] > 0.95 or ans[i] < 0.1 or ans[i] == ans[j]:
                    # print(ans[i], ans[j])
                    flag = 0
                    break
            if not flag:
                break
    # print('finally')
    # print(ans)
    return ans

def t5():
    flag = 0
    while not flag:
        flag = 1

        numbers = {3: 'три коробки', 4: 'четыре коробки', 5: 'пять коробок', 6: 'шесть коробок', 7: 'семь коробок'}
        boxes_1 = random.randrange(3, 7)  # выбрали количество коробок первой серии
        boxes_2 = abs(10 - boxes_1)

        white_1 = random.randint(2, 8)
        black_1 = abs(10 - white_1)

        white_2 = random.randint(2, 8)
        while white_2 == white_1:
            white_2 = random.randint(2, 8)
        black_2 = abs(10 - white_2)

        p_b1 = round(boxes_1 / 10, 5)  # вероятность выбора коробок первой серии
        p_b2 = round(boxes_2 / 10, 5)  # второй

        p_b1_a, p_b2_a = round(white_1 / 10, 5), round(white_2 / 10, 5)

        task = tasks[5].replace('четыре коробки', numbers[boxes_1], 1).replace('шесть коробок', numbers[boxes_2], 1).replace(' 3 ', ' $1$ ', 1).replace(' 7 ', ' $2$ ', 1).replace(' 8 ', ' $3$ ', 1).replace(' 2 ', ' $4$ ', 1)
        task = task.replace('$1$', f'{white_1}', 1).replace('$2$', f'{black_1}', 1).replace('$3$', f'{white_2}', 1).replace('$4$', f'{black_2}', 1)

        p = round((p_b1*p_b1_a) / (p_b1 * p_b1_a + p_b2 * p_b2_a), 7)
        w1 = round(p_b1 * p_b1_a + p_b2 * p_b2_a, 7)
        w2 = round((p_b2*p_b2_a) / (p_b1 * p_b1_a + p_b2 * p_b2_a), 7)
        w3 = round(p + .12, 7)

        ans = [task, p, w1, w2, w3]
        # print('iteration', ans)
        for i in range(1, len(ans)):
            for j in range(i + 1, len(ans)):
                if abs(ans[i] - ans[j]) <= 0.01 or ans[i] > 0.95 or ans[i] < 0.1 or ans[i] == ans[j] or len(str(ans[i])) > 6:
                    flag = 0
                    break
            if not flag:
                break
    # print('finally')
    # print(ans)
    return ans

def t6():
    flag = 0
    while not flag:
        flag = 1

        # Задание минимального и максимального значения для каждого числа
        min_value = 0.18
        max_value = 0.48
        # Генерация 4 случайных чисел в диапазоне от min_value до max_value с шагом 0.1
        numbers = [round(random.uniform(min_value, max_value), 2) for i in range(3)]
        numbers.append(round(1 - numbers[0]-numbers[1]-numbers[2], 2))
        while sum([abs(i) for i in numbers]) != 1:
            numbers = [round(random.uniform(min_value, max_value), 2) for i in range(3)]
            numbers.append(round(1 - numbers[0]-numbers[1]-numbers[2], 2))

        ans = [numbers, round(numbers[1]+numbers[2], 5), round(numbers[0]+numbers[1]+numbers[3], 5), numbers[0], numbers[1]]
        # print('iteration', ans[1:])
        for i in range(1, len(ans)):
            for j in range(i + 1, len(ans)):
                if abs(ans[i] - ans[j]) <= 0.03:
                    # print(ans[i], ans[j])
                    flag = 0
                    break
            if not flag:
                break
    # print('finally')
    # print(ans)
    return ans

def t7():
    flag = 0
    while not flag:
        flag = 1

        numbers = []
        numbers.append(round(random.uniform(0.12, 0.22), 2))
        numbers.append(round(random.uniform(0.12, 0.3), 2))
        numbers.append(round(0.6 - numbers[0] - numbers[1], 2))
        numbers.append(round(random.uniform(0.18, 0.3), 2))
        numbers.append(round(0.4 - numbers[3], 2))

        # print(numbers)

        p = (numbers[1], numbers[2], numbers[3])
        w1 = tuple(round(random.uniform(0.12, 0.22), 2) for i in range(3))
        w2 = tuple(round(random.uniform(0.12, 0.22), 2) for i in range(3))
        w3 = tuple(round(random.uniform(0.12, 0.22), 2) for i in range(3))

        ans = [numbers, p, w1, w2, w3]
        # print('iteration', ans)
        for i in range(1, len(ans)):
            for j in range(i + 1, len(ans)):
                if ans[i] == ans[j]:
                    flag = 0
    # print('finally')
    # print(ans)
    return ans

def t11():
    flag = 0
    while not flag:
        flag = 1

        p1 = round(random.uniform(0.09, 0.8), 2)
        p2 = round(1 - p1, 2)

        if p1 + p2 != 1:
            flag = 0
            continue

        # print(p1, p2)

        x1, x2 = 2,5
        M = round(p1 * x1 + p2 * x2, 4)

        p = p2
        w1 = p1
        w2 = round(random.uniform(0.09, 0.8), 2)
        w3 = round(random.uniform(0.09, 0.8), 2)

        ans = [(p1, p2, M), p, w1, w2, w3]
        # print('iteration', ans)
        for i in range(1, len(ans)):
            for j in range(i + 1, len(ans)):
                if abs(ans[i] - ans[j]) <= 0.01 or ans[i] > 0.95 or ans[i] < 0.1 or ans[i] == ans[j] or len(
                        str(ans[i])) > 6:
                    flag = 0
                    break
            if not flag:
                break
    # print('finally')
    # print(ans)
    return ans

def create_main_window():

    def validate_num_tests():
        try:
            num_tests = int(num_tests_entry.get())
            if num_tests < 1 or num_tests > 200:
                raise ValueError
        except ValueError:
            num_tests_entry.config(highlightbackground='red', highlightcolor='red')
            messagebox.showerror(title="Ошибка", message="Недопустимое значение. Введите число от 1 до 200.")
            return False
        else:
            num_tests_entry.config(highlightbackground='green', highlightcolor='green')
            return True


    root = tk.Tk()
    root.geometry("400x400")
    root.title("Генерация тестов")

    # Создаем заголовок
    title_label = tk.Label(root, text="Генерация тестов", font=("Helvetica", 20))
    title_label.pack(pady=10)

    # Создаем текстовое поле для ввода количества тестов
    num_tests_label = tk.Label(root, text="Количество тестов:")
    num_tests_label.pack(pady=10)

    # validate_cmd = (root.register(validate_num_tests), '%P')
    # num_tests_entry = tk.Entry(root, validate="key", validatecommand=validate_cmd)
    global num_tests_entry
    num_tests_entry = tk.Entry(root)
    num_tests_entry.pack()

    # Создаем стиль для кнопок
    button_style = ttk.Style()
    button_style.configure("Custom.TButton", background="gray", foreground="white", padding=10, font=("Helvetica", 12), borderwidth=0, focuscolor="none", focusthickness=0)
    button_style.map("Custom.TButton", background=[("active", "darkgray")], foreground=[("active", "white")])

    # Создаем кнопку для генерации всех тестов
    generate_tests_button = ttk.Button(root, text="Сгенерировать все тесты", style="Custom.TButton", command=lambda: (generate_tests(int(num_tests_entry.get())), generate_teor_tests(int(num_tests_entry.get())), messagebox.showinfo(title="Успешно", message=f"Сгенерировано тестов: {num_tests_entry.get()}")))
    generate_tests_button.bind("<ButtonPress>", lambda event: validate_num_tests())
    generate_tests_button.pack(pady=10)

    # Создаем кнопку для генерации практических тестов
    generate_tests_button = ttk.Button(root, text="Сгенерировать практические тесты", style="Custom.TButton", command=lambda: (generate_tests(int(num_tests_entry.get())), messagebox.showinfo(title="Успешно", message=f"Сгенерировано практических тестов: {num_tests_entry.get()}")))
    generate_tests_button.bind("<ButtonPress>", lambda event: validate_num_tests())
    generate_tests_button.pack(pady=10)

    # Создаем кнопку для генерации теоретических тестов
    generate_tests_button = ttk.Button(root, text="Сгенерировать теоретические тесты", style="Custom.TButton", command=lambda: (generate_teor_tests(int(num_tests_entry.get())), messagebox.showinfo(title="Успешно", message=f"Сгенерировано теоретических тестов: {num_tests_entry.get()}")))
    generate_tests_button.bind("<ButtonPress>", lambda event: validate_num_tests())
    generate_tests_button.pack(pady=10)



    # # НЕ создаем кнопку для скачивания файла
    # download_file_button = ttk.Button(root, text="Скачать примеры", style="Custom.TButton")
    # download_file_button.pack(pady=10)

    def about():
        messagebox.showinfo(title="О программе",
                           message="Версия 0.8\n\nАвторы:\nКолычев Егор\nКорнилов Кирилл\nПолевая Полина",
                           detail="© MIT License. 2023.")


    menu = tk.Menu(root)
    root.config(menu=menu)

    help_menu = tk.Menu(menu, tearoff=False)
    menu.add_cascade(label="Помощь", menu=help_menu)
    help_menu.add_command(label="О программе", command=about)

    # Показываем главное окно
    root.mainloop()

def generate_tests(num_tests):

    def table_style():
        # table.style = 'Table Grid'
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

        table_font = table.style.font
        table_font.name = 'Times New Roman'
        table_font.size = docx.shared.Pt(16)

        for row in table.rows:
            for cell in row.cells:
                cell_font = cell.paragraphs[0].style.font
                cell_font.name = 'Times New Roman'
                cell_font.size = docx.shared.Pt(16)

    # num_tests = int(input('Количество тестов для генерации: '))

    document = docx.Document()  # документ с практическими тестами

    # задание стиля для header
    style_header = document.styles.add_style('f_header', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
    style_header.font.name = 'Times New Roman'
    style_header.font.size = docx.shared.Pt(16)
    style_header.font.italic = True

    # задание стиля для заданий
    style_task = document.styles.add_style('f_tasks', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
    style_task.font.name = 'Times New Roman'
    style_task.font.size = docx.shared.Pt(16)

    for i in range(1, num_tests+1):
        answers[i] = dict()
        # добавление параграфа с вариантом
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Вариант 4 (№{i})')
        run.style = style_header
        run.font.bold = True
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # добавление блока с фамилией и группой
        paragraph = document.add_paragraph()
        run = paragraph.add_run('\nФамилия ________________________ Группа __________')
        run.style = style_header
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # блок заданий

        # задание 1
        paragraph = document.add_paragraph()
        run = paragraph.add_run('1. ')
        run.style = style_task
        run.bold = True

        task = t1()
        run = paragraph.add_run(task[0])
        run.style = style_task

        task_ans = task[1:]
        p_find = task_ans[0]
        random.shuffle(task_ans)

        # print(p_find, task_ans, chr(task_ans.index(p_find)+1072).capitalize())
        answers[i][1] = chr(task_ans.index(p_find)+1072).capitalize()
        # print(answers)



        table = document.add_table(rows=1, cols=4)
        table_style()

        row_cells = table.rows[0].cells
        row_cells[0].text = f"а) {task_ans[0]};"
        row_cells[1].text = f"б) {task_ans[1]};"
        row_cells[2].text = f"в) {task_ans[2]};"
        row_cells[3].text = f"г) {task_ans[3]}."

        # задание 2
        paragraph = document.add_paragraph()
        run = paragraph.add_run('2. ')
        run.style = style_task
        run.bold = True

        run = paragraph.add_run(tasks[2])
        run.style = style_task

        task_ans = [f'1/2\u03C0', f'2/\u03C0', f'\u03C0/36', f'\u221A3/4']
        p_find = task_ans[0]
        random.shuffle(task_ans)
        answers[i][2] = chr(task_ans.index(p_find) + 1072).capitalize()

        table = document.add_table(rows=1, cols=4)
        table_style()
        row_cells = table.rows[0].cells
        row_cells[0].text = f"а) {task_ans[0]};"
        row_cells[1].text = f"б) {task_ans[1]};"
        row_cells[2].text = f"в) {task_ans[2]};"
        row_cells[3].text = f"г) {task_ans[3]}."

        # задание 3
        paragraph = document.add_paragraph()
        run = paragraph.add_run('3. ')
        run.style = style_task
        run.bold = True

        task = t3()
        run = paragraph.add_run(task[0])
        run.style = style_task

        task_ans = task[1:]
        p_find = task_ans[0]
        random.shuffle(task_ans)
        answers[i][3] = chr(task_ans.index(p_find) + 1072).capitalize()

        table = document.add_table(rows=1, cols=4)
        table_style()
        row_cells = table.rows[0].cells
        row_cells[0].text = f"а) {task_ans[0]};"
        row_cells[1].text = f"б) {task_ans[1]};"
        row_cells[2].text = f"в) {task_ans[2]};"
        row_cells[3].text = f"г) {task_ans[3]}."

        # задание 4
        paragraph = document.add_paragraph()
        run = paragraph.add_run('4. ')
        run.style = style_task
        run.bold = True

        task = t4()
        run = paragraph.add_run(task[0])
        run.style = style_task
        task_ans = task[1:]
        p_find = task_ans[0]
        random.shuffle(task_ans)
        answers[i][4] = chr(task_ans.index(p_find) + 1072).capitalize()

        table = document.add_table(rows=1, cols=4)
        table_style()
        row_cells = table.rows[0].cells
        row_cells[0].text = f"а) {task_ans[0]};"
        row_cells[1].text = f"б) {task_ans[1]};"
        row_cells[2].text = f"в) {task_ans[2]};"
        row_cells[3].text = f"г) {task_ans[3]}."

        # задание 5
        paragraph = document.add_paragraph()
        run = paragraph.add_run('5. ')
        run.style = style_task
        run.bold = True

        task = t5()
        run = paragraph.add_run(task[0])
        run.style = style_task
        task_ans = task[1:]
        p_find = task_ans[0]
        random.shuffle(task_ans)
        answers[i][5] = chr(task_ans.index(p_find) + 1072).capitalize()

        table = document.add_table(rows=1, cols=4)
        table_style()
        row_cells = table.rows[0].cells
        row_cells[0].text = f"а) {task_ans[0]};"
        row_cells[1].text = f"б) {task_ans[1]};"
        row_cells[2].text = f"в) {task_ans[2]};"
        row_cells[3].text = f"г) {task_ans[3]}."

        # задание 6
        paragraph = document.add_paragraph()
        run = paragraph.add_run('6. ')
        run.style = style_task
        run.bold = True

        run = paragraph.add_run(tasks[6][0])
        run.style = style_task

        task_ans = t6()

        table = document.add_table(rows=2, cols=5)
        table.style = 'Table Grid'

        row_cells = table.rows[0].cells
        p = row_cells[0].paragraphs[0]
        run = p.add_run()
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = 'x\u1D62'
        run._r.append(t)
        row_cells[1].text = f"1"
        row_cells[2].text = f"2"
        row_cells[3].text = f"4"
        row_cells[4].text = f"6"

        row_cells = table.rows[1].cells
        p = row_cells[0].paragraphs[0]
        run = p.add_run()
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = 'p\u1D62'
        run._r.append(t)
        row_cells[1].text = f"{task_ans[0][0]}"
        row_cells[2].text = f"{task_ans[0][1]}"
        row_cells[3].text = f"{task_ans[0][2]}"
        row_cells[4].text = f"{task_ans[0][3]}"

        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER  # располагаем таблицу по центру
        for row in table.rows:
            for cell in row.cells:
                cell.width = docx.shared.Inches(0.8)
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # устанавливаем выравнивание ячеек по центру по вертикали
                    for run in paragraph.runs:
                        run.font.bold = False  # убираем жирный шрифт

        paragraph = document.add_paragraph()
        task2 = tasks[6][1].split('P(<)')
        run = paragraph.add_run("\n   " + task2[0])
        run.style = style_task
        run = paragraph.add_run('P')
        run.font.italic = True
        run = paragraph.add_run(f"(1 < ")
        run.font.italic = False
        run = paragraph.add_run("X")
        run.font.italic = True
        run = paragraph.add_run(f' \u2264 4)')
        run.font.italic = False
        run = paragraph.add_run(task2[1])

        task_ans = task_ans[1:]
        p_find = task_ans[0]
        random.shuffle(task_ans)
        answers[i][6] = chr(task_ans.index(p_find) + 1072).capitalize()

        table = document.add_table(rows=1, cols=4)
        table_style()
        row_cells = table.rows[0].cells
        row_cells[0].text = f"а) {task_ans[0]};"
        row_cells[1].text = f"б) {task_ans[1]};"
        row_cells[2].text = f"в) {task_ans[2]};"
        row_cells[3].text = f"г) {task_ans[3]}."

        # задание 7
        paragraph = document.add_paragraph()
        run = paragraph.add_run('7. ')
        run.style = style_task
        run.bold = True

        run = paragraph.add_run(tasks[7][0])
        run.style = style_task

        task_ans = t7()

        table = document.add_table(rows=2, cols=6)
        table.style = 'Table Grid'

        row_cells = table.rows[0].cells
        p = row_cells[0].paragraphs[0]
        run = p.add_run()
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = 'x\u1D62'
        run._r.append(t)
        row_cells[1].text = f"1"
        row_cells[2].text = f"3"
        row_cells[3].text = f"5"
        row_cells[4].text = f"7"
        row_cells[5].text = f"9"

        row_cells = table.rows[1].cells
        p = row_cells[0].paragraphs[0]
        run = p.add_run()
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = 'p\u1D62'
        run._r.append(t)
        row_cells[1].text = f"{task_ans[0][0]}"
        row_cells[2].text = f"{task_ans[0][1]}"
        row_cells[3].text = f"{task_ans[0][2]}"
        row_cells[4].text = f"{task_ans[0][3]}"
        row_cells[5].text = f"{task_ans[0][4]}"

        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER  # располагаем таблицу по центру
        for row in table.rows:
            for cell in row.cells:
                cell.width = docx.shared.Inches(0.8)
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # устанавливаем выравнивание ячеек по центру по вертикали
                    for run in paragraph.runs:
                        run.font.bold = False  # убираем жирный шрифт

        paragraph = document.add_paragraph()
        task2 = tasks[7][1].split('P(<)')
        run = paragraph.add_run("\n   " + task2[0])
        run.style = style_task
        run = paragraph.add_run('P')
        run.font.italic = True
        run = paragraph.add_run(f"(1 \u2264 ")
        run.font.italic = False
        run = paragraph.add_run("X")
        run.font.italic = True
        run = paragraph.add_run(f' \u2264 4)')
        run.font.italic = False
        run = paragraph.add_run(task2[1])

        task_ans = task_ans[1:]
        p_find = task_ans[0]
        random.shuffle(task_ans)
        answers[i][7] = chr(task_ans.index(p_find) + 1072).capitalize()

        table = document.add_table(rows=2, cols=2)
        table_style()
        row_cells = table.rows[0].cells
        row_cells[0].text = f"а) a = {task_ans[0][0]}; b = {task_ans[0][1]}; c = {task_ans[0][2]};"
        row_cells[1].text = f"б) a = {task_ans[1][0]}; b = {task_ans[1][1]}; c = {task_ans[1][2]};"
        row_cells = table.rows[1].cells
        row_cells[0].text = f"в) a = {task_ans[2][0]}; b = {task_ans[2][1]}; c = {task_ans[2][2]};"
        row_cells[1].text = f"г) a = {task_ans[3][0]}; b = {task_ans[3][1]}; c = {task_ans[3][2]}."

        # задание 8
        paragraph = document.add_paragraph()
        task = "Непрерывная случайная величина X задана плотностью распределения вероятностей:\t"
        run = paragraph.add_run('8. ')
        run.bold = True
        run = paragraph.add_run(task)
        run.style = style_task
        run.bold = False
        answers[i][8] = printTask8(document)
        # paragraph = document.add_paragraph()

        # задание 9
        paragraph = document.add_paragraph()
        task = "Непрерывная случайная величина X задана плотностью распределения вероятностей:\t"
        run = paragraph.add_run('9. ')
        run.bold = True
        run = paragraph.add_run(task)
        run.bold = False
        run.style = style_task
        answers[i][9] = printTask9(document)
        # paragraph = document.add_paragraph()

        # задание 10
        paragraph = document.add_paragraph()
        task = "Непрерывная случайная величина X задана функцией распределения вероятностей:\t"
        run = paragraph.add_run('10. ')
        run.bold = True
        run = paragraph.add_run(task)
        run.bold = False
        run.style = style_task
        answers[i][10] = printTask10(document)
        # paragraph = document.add_paragraph()

        # задание 11
        paragraph = document.add_paragraph()
        run = paragraph.add_run('11. ')
        run.style = style_task
        run.bold = True

        run = paragraph.add_run(f"{tasks[11][0]}\t")
        run.style = style_task

        task_ans = t11()

        table = document.add_table(rows=2, cols=3)
        table.style = 'Table Grid'

        row_cells = table.rows[0].cells
        p = row_cells[0].paragraphs[0]
        run = p.add_run()
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = 'x\u1D62'
        run._r.append(t)
        row_cells[1].text = f"2"
        row_cells[2].text = f"5"

        row_cells = table.rows[1].cells
        p = row_cells[0].paragraphs[0]
        run = p.add_run()
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = 'p\u1D62'
        run._r.append(t)
        row_cells[1].text = f"\u0070\u2081"
        row_cells[2].text = f"\u0070\u2082"

        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER  # располагаем таблицу по центру
        for row in table.rows:
            for cell in row.cells:
                cell.width = docx.shared.Inches(0.8)
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # устанавливаем выравнивание ячеек по центру по вертикали
                    for run in paragraph.runs:
                        run.font.bold = False  # убираем жирный шрифт

        paragraph = document.add_paragraph()
        # print()
        task2 = tasks[11][1].split('p2')
        task2[0] = task2[0].replace('4,7', str(task_ans[0][2]))
        run = paragraph.add_run("\n   " + task2[0])
        run.style = style_task
        run = paragraph.add_run("\u0070\u2082")
        run = paragraph.add_run(task2[1])

        task_ans = task_ans[1:]
        p_find = task_ans[0]
        random.shuffle(task_ans)
        answers[i][11] = chr(task_ans.index(p_find) + 1072).capitalize()

        table = document.add_table(rows=1, cols=4)
        table_style()
        row_cells = table.rows[0].cells
        row_cells[0].text = f"а) {task_ans[0]};"
        row_cells[1].text = f"б) {task_ans[1]};"
        row_cells[2].text = f"в) {task_ans[2]};"
        row_cells[3].text = f"г) {task_ans[3]}."

        # paragraph = document.add_paragraph()
        # printToMathml(paragraph, "A_i = \lambda")

        # задание 12
        paragraph = document.add_paragraph()
        task = "Непрерывная случайная величина X задана плотностью распределения вероятностей:\t"
        run = paragraph.add_run('12. ')
        run.bold = True
        run = paragraph.add_run(task)
        run.bold = False
        run.style = style_task
        answers[i][12] = printTask12(document)
        # paragraph = document.add_paragraph()

        # задание 13
        M = random.randint(1, 31)  # Создаем мат ожидание
        D = (random.randint(1, 10)) ** 2  # Создаем дисперсию

        task = 'Случайная величина X распределена нормально с математическим ожиданием M(X) = ' + str(
            M) + ' и дисперсией D(X) = ' + str(D) + '. Тогда ее плотность распределения вероятностей имеет вид:'

        paragraph = document.add_paragraph()

        run = paragraph.add_run('13. ')
        ran = paragraph.add_run(task)
        run.style = style_task
        run.bold = True
        answers[i][13] = printTask13(document, M, D)
        paragraph = document.add_paragraph()

        # перенос страницы
        if i != num_tests:
            document.add_page_break()
        else:
            print('Тесты практика сгенерены')

    document.save('text.docx')
    # document.save(os.path.join(save_folder, 'text.docx'))generate_answers

    generate_answers(num_tests)

def generate_answers(num_tests):
    # print(answers)
    doc_ans = docx.Document()  # документ с ответами на практический тест


    #Добавлен заголовок
    title=doc_ans.add_paragraph(f'Ответы для тестов «Варианты (1-{num_tests})»')
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    font = title.style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(16)
    font.bold = True
    run=title.add_run()
    run.add_break(docx.enum.text.WD_BREAK.LINE)


    table = doc_ans.add_table(rows=1 + num_tests, cols=14,style='Table Grid')
    # # Задаем ширину каждого столбца
    # column_widths = [docx.shared.Inches(0.6) for i in range(13)]
    # column_widths[0] = docx.shared.Inches(2)
    #
    # # Устанавливаем ширину каждого столбца
    # for i, width in enumerate(column_widths):
    #     table.columns[i].width = width

    # Добавляем данные в каждую ячейку таблицы
    row_cells = table.rows[0].cells
    row_cells[0].text = 'B\№'
    for i in range(1, 13 + 1):
        row_cells[i].text = str(i)

    for row in table.rows:
        for cell in row.cells:
            cell.width = docx.shared.Inches(0.5)
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # устанавливаем выравнивание ячеек по центру по вертикали
                for run in paragraph.runs:
                    run.font.bold = True  # убираем жирный шрифт

    for i in range(1, num_tests + 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)
        for j in range(1, 13 + 1):
            row_cells[j].text = answers[i][j]

    # Выравниваем текст в ячейках по центру
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = docx.shared.Pt(14)
                    run.font.name = 'Times New Roman'
                    run.font.bold = False
                    run.font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0x00)


    # Делаем жирными заголовки
    cells = table.column_cells(0)
    for cell in cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            paragraph.alignment =docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
    cells = table.row_cells(0)
    for cell in cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True

    # Выравниваем таблицу по центру страницы
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    # Сохраняем документ
    doc_ans.save('table.docx')

if __name__ == '__main__':
    # # Путь к папке с проектом
    # project_folder = os.path.dirname(os.path.abspath(__file__))
    #
    # # Путь к папке для сохранения файлов
    # save_folder = os.path.join(project_folder, "сформированные файлы")
    #
    # # Создаем папку, если она еще не существует
    # if not os.path.exists(save_folder):
    #     os.makedirs(save_folder)

    create_main_window()